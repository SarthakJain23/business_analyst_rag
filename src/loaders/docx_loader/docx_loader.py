from pathlib import Path
from typing import List

import docx

from src.loaders.base.base import BaseLoader, RawDocument
from src.utils.logger import get_logger

logger = get_logger("docx_loader")


class DocxLoader(BaseLoader):
    """Loader for Microsoft Word (.docx) documents."""

    def load(self, file_path: Path, file_hash: str) -> List[RawDocument]:
        documents: List[RawDocument] = []
        try:
            doc = docx.Document(str(file_path))
            full_text = []

            for p in doc.paragraphs:
                if p.text.strip():
                    full_text.append(p.text.strip())

            for table_idx, table in enumerate(doc.tables):
                table_lines = [f"\n### Table {table_idx + 1}"]
                for row in table.rows:
                    row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    table_lines.append("| " + " | ".join(row_cells) + " |")
                full_text.append("\n".join(table_lines))

            combined_content = "\n\n".join(full_text)
            if combined_content.strip():
                metadata = {
                    "source_file": str(file_path.resolve()),
                    "file_name": file_path.name,
                    "file_type": ".docx",
                    "file_hash": file_hash,
                    "page_or_section": "Document Body",
                }
                documents.append(RawDocument(content=combined_content, metadata=metadata))

            logger.info(f"Successfully loaded DOCX document: {file_path.name}")
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
            raise e

        return documents
